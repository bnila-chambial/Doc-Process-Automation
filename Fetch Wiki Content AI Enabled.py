import requests
from bs4 import BeautifulSoup
from docx import Document

CONFLUENCE_URL = "https://wiki.landisgyr.net"
SPACE_KEY = "US"

SECTION_HEADER = "New Features"
WORD_DOC_PATH = "C:\\Users\\PawarSar\\Documents\\Test Template.docx"
OUTPUT_DOC_PATH = "C:\\Users\\PawarSar\\Documents\\Updated Template.docx"
PARAPHRASED_DOC_PATH = "C:\\Users\\PawarSar\\Documents\\Paraphrased Template.docx"
PLACEHOLDER_TEXT = "[[New Features22]]"

def get_password(prompt='Password: '):
    import msvcrt
    print(prompt, end='', flush=True)
    pw = ''
    while True:
        ch = msvcrt.getch()
        if ch in {b'\r', b'\n'}:
            print('')
            break
        elif ch == b'\x08':  # Backspace
            if len(pw) > 0:
                pw = pw[:-1]
                print('\b \b', end='', flush=True)
        elif ch == b'\x03':  # Ctrl+C
            raise KeyboardInterrupt
        else:
            pw += ch.decode('utf-8', errors='ignore')
            print('*', end='', flush=True)
    return pw

def paraphrase_with_ollama(text, model="llama2"):
    """
    Sends the text to the local Ollama server for paraphrasing using the specified model.
    """
    url = "http://localhost:11434/api/generate"
    prompt = f"Paraphrase the following text in clear, professional English:\n\n{text}"
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False
    }
    response = requests.post(url, json=payload)
    response.raise_for_status()
    result = response.json()
    return result.get("response", "").strip()

def paraphrase_docx(input_path, output_path):
    doc = Document(input_path)
    for para in doc.paragraphs:
        if para.text.strip():
            try:
                para.text = paraphrase_with_ollama(para.text)
            except Exception as e:
                print(f"Error paraphrasing: {e}")
    doc.save(output_path)
    print(f"✅ Paraphrased document saved as '{output_path}'")

username = input("Enter your Confluence username: ")
password = get_password("Enter your Confluence password: ")
get_rn_version = input ("Enter the release notes version you want to create, e.g., '8.6', '8.7', etc.:---->")

PAGE_TITLE = f"TS {get_rn_version} Release Notes"
# Step 1: Search for the page to get its ID
search_url = f"{CONFLUENCE_URL}/rest/api/content"
params = {'title': PAGE_TITLE, 'spaceKey': SPACE_KEY}
response = requests.get(search_url, params=params, auth=(username, password))
response.raise_for_status()
data = response.json()

if not data['results']:
    raise ValueError("Page not found.")

page_id = data['results'][0]['id']

# Step 2: Fetch the page content with body.storage
page_url = f"{CONFLUENCE_URL}/rest/api/content/{page_id}?expand=body.storage"
response = requests.get(page_url, auth=(username, password))
response.raise_for_status()
page_data = response.json()
html_content = page_data['body']['storage']['value']

# Parse HTML
soup = BeautifulSoup(html_content, 'html.parser')

# Find placeholder paragraph index
doc = Document(WORD_DOC_PATH)
for i, para in enumerate(doc.paragraphs):
    if PLACEHOLDER_TEXT in para.text:
        target_index = i
        break
else:
    raise ValueError("Placeholder not found in document.")

# Remove placeholder paragraph
p = doc.paragraphs[target_index]
p._element.getparent().remove(p._element)

# Get all the content from the h3 section

def get_h3_section_content(soup, keyword):
    """
    Returns all content (as text) directly under the h3 heading matching `keyword`,
    until the next h3/h2/h1 or end of document.
    """
    content = []
    found = False
    for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol']):
        if not found:
            if elem.name == 'h3' and elem.get_text(strip=True) == keyword:
                found = True
            continue
        # Stop if we hit another heading of same or higher level
        if elem.name in ['h1', 'h2', 'h3']:
            break
        if elem.name in ['p', 'h4']:
            content.append(elem.get_text(strip=True))
        elif elem.name in ['ul', 'ol']:
            for li in elem.find_all('li', recursive=False):
                content.append(li.get_text(strip=True))
    return '\n'.join(content)

all_content = get_h3_section_content(soup, SECTION_HEADER)

# Paraphrase the section content before inserting

paraphrased_content = paraphrase_with_ollama(all_content) # !!! remove this line to paraphrase the whole document. !!!

# Insert as a single paragraph at the target location
def insert_content_at(doc, index, content):
    para = doc.add_paragraph(content)
    para_element = para._element
    body = doc._body._element
    body.remove(para_element)
    body.insert(index, para_element)

insert_content_at(doc, target_index + 1, all_content)

doc.save(OUTPUT_DOC_PATH)
print(f"✅ Section inserted into '{OUTPUT_DOC_PATH}' successfully.")

# Paraphrase the updated document using ollama
paraphrase_docx(OUTPUT_DOC_PATH, PARAPHRASED_DOC_PATH)